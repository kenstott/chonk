# Copyright (c) 2025 Kenneth Stott. MIT License.
# Canary: 1d16e713-d096-4779-b8fb-b68927321157
#
# NOTICE: Use of this software for training artificial intelligence or
# machine learning models is strictly prohibited without explicit written
# permission from the copyright holder.

"""DOCX text extractor using python-docx."""

from __future__ import annotations

from io import BytesIO
from typing import TYPE_CHECKING

if TYPE_CHECKING:
    from chonk.models import DocumentChunk

try:
    import docx  # noqa: F401

    _DOCX_AVAILABLE = True
except ImportError:
    _DOCX_AVAILABLE = False


def _style_name(para) -> str:
    return para.style.name if para.style is not None else ""


def _extract_docx_content(doc) -> str:
    paragraphs = []

    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            name = _style_name(para)
            if name.startswith("Heading"):
                level = name.replace("Heading ", "")
                try:
                    level_num = int(level)
                    paragraphs.append(f"{'#' * level_num} {text}")
                except ValueError:
                    paragraphs.append(text)
            else:
                paragraphs.append(text)

    for table in doc.tables:
        table_rows = []
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            table_rows.append(" | ".join(cells))
        if table_rows:
            paragraphs.append("\n".join(table_rows))

    return "\n\n".join(paragraphs)


class DocxExtractor:
    def can_handle(self, doc_type: str) -> bool:
        return doc_type == "docx"

    def extract(self, data: bytes, source_path: str | None = None) -> str:
        if not _DOCX_AVAILABLE:
            raise ImportError(f"pip install chonk[docx] (loading {source_path or 'unknown'})")
        import docx as _docx

        doc = _docx.Document(BytesIO(data))
        return _extract_docx_content(doc)

    def annotate(self, chunks: list[DocumentChunk], data: bytes, source_path: str | None = None) -> list[DocumentChunk]:
        if not _DOCX_AVAILABLE:
            raise ImportError(f"pip install chonk[docx] (loading {source_path or 'unknown'})")
        import docx as _docx

        doc = _docx.Document(BytesIO(data))

        para_records: list[tuple[int, str, list[str]]] = []
        heading_stack: list[str] = []
        for orig_idx, para in enumerate(doc.paragraphs):
            text = para.text.strip()
            if not text:
                continue
            name = _style_name(para)
            if name.startswith("Heading"):
                try:
                    level = int(name.replace("Heading ", ""))
                    heading_stack = heading_stack[: level - 1] + [text]
                except ValueError:
                    pass
            para_records.append((orig_idx, text, list(heading_stack)))

        for chunk in chunks:
            content = chunk.content
            start_idx: int | None = None
            end_idx: int | None = None
            section_at_start: list[str] = []
            for orig_idx, text, section in para_records:
                if len(text) >= 10 and text in content:
                    if start_idx is None:
                        start_idx = orig_idx
                        section_at_start = section
                    end_idx = orig_idx
            if start_idx is not None:
                chunk.source_detail = {
                    "paragraph_start": start_idx,
                    "paragraph_end": end_idx,
                    "section": section_at_start,
                }

        return chunks
